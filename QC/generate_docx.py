import docx
from docx.shared import Pt
import re
import sys

md_path = r"c:\Users\Faiz\.gemini\antigravity-ide\brain\a1fd7d1d-1481-45b8-9ecd-8158b02cb648\qc_checklist.md"
doc_path = r"C:\Users\Faiz\OneDrive\Dokumen\UNPAM\Semester 6\KP\PEMBUATAN WEBSITE ABSENSI SISWA DAN PENENTUAN SISWA TERBAIK\QC\Laporan_Hasil_QC.docx"

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def add_paragraph(text, style='Normal'):
    # Basic markdown parsing for bold
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # check for italic
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = p.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    p.add_run(subpart)
    return p

try:
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:].replace('*', ''), 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].replace('*', ''), 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].replace('*', ''), 2)
        elif line.startswith('- [x]') or line.startswith('- [ ]'):
            is_checked = line.startswith('- [x]')
            symbol = "✅" if is_checked else "❌"
            content = line[5:].strip()
            add_paragraph(f"{symbol} {content}")
        elif line.startswith('- '):
            add_paragraph("• " + line[2:].strip())
        elif line.startswith('```') or line.startswith('graph') or line.startswith('Fase'):
            pass # Skip mermaid code blocks
        else:
            add_paragraph(line)

    doc.save(doc_path)
    print("Berhasil membuat file DOCX di", doc_path)
except Exception as e:
    print("Error:", str(e))
    sys.exit(1)
